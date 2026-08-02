"""Document text extraction adapters (PDF and Word .docx)."""

from __future__ import annotations

from abc import ABC, abstractmethod
from pathlib import Path

import pdfplumber
from docx import Document

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
SUPPORTED_CONTENT_TYPES = {PDF_CONTENT_TYPE, DOCX_CONTENT_TYPE}


class DocumentTextExtractor(ABC):
    @abstractmethod
    def extract(self, path: Path) -> str:
        raise NotImplementedError


class PdfplumberExtractor(DocumentTextExtractor):
    def extract(self, path: Path) -> str:
        chunks: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    chunks.append(text)
        return "\n\n".join(chunks).strip()


class DocxExtractor(DocumentTextExtractor):
    """Extract plain text from modern Word documents (.docx) via python-docx."""

    def extract(self, path: Path) -> str:
        doc = Document(path)
        chunks: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))

        return "\n\n".join(chunks).strip()


class CompositeDocumentExtractor(DocumentTextExtractor):
    """Route extraction to PDF or DOCX adapter based on file extension."""

    def __init__(
        self,
        *,
        pdf_extractor: DocumentTextExtractor | None = None,
        docx_extractor: DocumentTextExtractor | None = None,
    ) -> None:
        self._pdf = pdf_extractor or PdfplumberExtractor()
        self._docx = docx_extractor or DocxExtractor()

    def extract(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._pdf.extract(path)
        if suffix == ".docx":
            return self._docx.extract(path)
        raise ValueError(f"Unsupported document type: {suffix or path.name}")


def resolve_upload_format(filename: str, content_type: str | None) -> tuple[str, str]:
    """
    Validate upload and return (extension_with_dot, canonical_content_type).
    Legacy binary Word (.doc) is not supported — only .docx.
    """
    name = (filename or "").lower()
    ct = (content_type or "").lower().split(";")[0].strip()

    if name.endswith(".doc") and not name.endswith(".docx"):
        raise ValueError(
            "Legacy Word (.doc) is not supported. Please upload .docx or PDF."
        )

    if name.endswith(".pdf") or ct == PDF_CONTENT_TYPE:
        return ".pdf", PDF_CONTENT_TYPE
    if name.endswith(".docx") or ct == DOCX_CONTENT_TYPE:
        return ".docx", DOCX_CONTENT_TYPE

    raise ValueError("Only PDF and Word (.docx) files are supported")
